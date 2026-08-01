# -*- coding: utf-8 -*-
"""Convert the eval/acceptance markdown report to .docx (stdlib + python-docx)."""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

MD = r"C:\Users\admin\Documents\Codex\2026-08-01\new-chat\outputs\Eval与项目验收报告_风险人工介入评估器.md"
OUT = r"C:\Users\admin\Documents\Codex\2026-08-01\new-chat\outputs\Eval与项目验收报告_风险人工介入评估器.docx"

def add_runs(par, text, base_bold=False):
    # handle **bold** and `code`
    tokens = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for t in tokens:
        if not t:
            continue
        if t.startswith("**") and t.endswith("**"):
            r = par.add_run(t[2:-2]); r.bold = True
        elif t.startswith("`") and t.endswith("`"):
            r = par.add_run(t[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9)
        else:
            par.add_run(t)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Microsoft YaHei"
style.font.size = Pt(10.5)

lines = open(MD, encoding="utf-8").read().splitlines()
i = 0
tbl_rows = []
in_code = False
while i < len(lines):
    line = lines[i].rstrip()
    if line.startswith("```"):
        in_code = not in_code
        i += 1
        continue
    if in_code:
        p = doc.add_paragraph(); r = p.add_run(line); r.font.name = "Consolas"; r.font.size = Pt(9)
        i += 1
        continue
    if not line.strip():
        i += 1
        continue
    if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:\-|]+\|$", lines[i+1].strip()):
        # table
        header = [c.strip() for c in line.strip("|").split("|")]
        rows = []
        j = i + 2
        while j < len(lines) and lines[j].strip().startswith("|"):
            rows.append([c.strip() for c in lines[j].strip().strip("|").split("|")])
            j += 1
        table = doc.add_table(rows=1 + len(rows), cols=len(header))
        table.style = "Table Grid"
        for k, h in enumerate(header):
            cell = table.rows[0].cells[k]
            cell.text = ""
            add_runs(cell.paragraphs[0], h, base_bold=True)
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(9)
        for r_i, row in enumerate(rows):
            for k in range(len(header)):
                val = row[k] if k < len(row) else ""
                cell = table.rows[r_i + 1].cells[k]
                cell.text = ""
                add_runs(cell.paragraphs[0], val)
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(9)
        doc.add_paragraph()
        i = j
        continue
    if line.startswith("### "):
        doc.add_heading(line[4:], level=3); i += 1; continue
    if line.startswith("## "):
        doc.add_heading(line[3:], level=2); i += 1; continue
    if line.startswith("# "):
        doc.add_heading(line[2:], level=1); i += 1; continue
    if line == "---":
        i += 1; continue
    m = re.match(r"^\s*[-*] \[[ xX]\]\s+(.*)$", line)
    if m:
        p = doc.add_paragraph(style="List Bullet")
        add_runs(p, ("✅ " if "[x]" in line.lower() else "⬜ ") + m.group(1))
        i += 1; continue
    m = re.match(r"^\s*[-*]\s+(.*)$", line)
    if m:
        p = doc.add_paragraph(style="List Bullet")
        add_runs(p, m.group(1))
        i += 1; continue
    m = re.match(r"^\s*\d+\.\s+(.*)$", line)
    if m:
        p = doc.add_paragraph(style="List Number")
        add_runs(p, m.group(1))
        i += 1; continue
    p = doc.add_paragraph()
    add_runs(p, line)
    i += 1

doc.save(OUT)
print("saved", OUT)
